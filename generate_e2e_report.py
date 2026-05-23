"""Generate E2E test report as .docx with embedded screenshots."""
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os

doc = Document()

# Title
title = doc.add_heading("E2E Test Report — TesAxrail App", level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Metadata
doc.add_paragraph(f"Device: Pixel 7 (Android 17, emulator-5554)")
doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
doc.add_paragraph("Results: 5/5 passed, 0 failed")
doc.add_paragraph("")

# Test Results Table
doc.add_heading("Test Results", level=2)

tests = [
    ("1", "Greeting Flow", '"Hello, Kiro!" appears', '"Hello, Kiro!" appeared', "✅ PASS", "e2e-android-01-greeting-flow-pass.png"),
    ("2", "Counter Increment (x3)", "Counter: 3", "Counter: 3", "✅ PASS", "e2e-android-02-counter-increment-pass.png"),
    ("3", "Counter Decrement", "Counter: 2", "Counter: 2", "✅ PASS", "e2e-android-03-counter-decrement-pass.png"),
    ("4", "Counter Reset", "Counter: 0", "Counter: 0", "✅ PASS", "e2e-android-04-counter-reset-pass.png"),
    ("5", "Negative Counter", "Counter: -1", "Counter: -1", "✅ PASS", "e2e-android-05-negative-counter-pass.png"),
]

table = doc.add_table(rows=len(tests) + 1, cols=5)
table.style = "Table Grid"

headers = ["#", "Test Case", "Expected", "Actual", "Status"]
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    cell.paragraphs[0].runs[0].bold = True

for row_idx, test in enumerate(tests, start=1):
    table.rows[row_idx].cells[0].text = test[0]
    table.rows[row_idx].cells[1].text = test[1]
    table.rows[row_idx].cells[2].text = test[2]
    table.rows[row_idx].cells[3].text = test[3]
    table.rows[row_idx].cells[4].text = test[4]

doc.add_paragraph("")

# Screenshots
doc.add_heading("Screenshots", level=2)

screenshots_dir = "/Users/yobel/AndroidStudioProjects/TesAxrail/e2e-screenshots"

# Initial state
doc.add_paragraph("Initial State (fresh app launch):")
initial = os.path.join(screenshots_dir, "e2e-android-00-initial-state.png")
if os.path.exists(initial):
    doc.add_picture(initial, width=Inches(2.5))
doc.add_paragraph("")

# Each test screenshot
for test in tests:
    doc.add_paragraph(f"Test #{test[0]}: {test[1]} — {test[4]}")
    screenshot_path = os.path.join(screenshots_dir, test[5])
    if os.path.exists(screenshot_path):
        doc.add_picture(screenshot_path, width=Inches(2.5))
    doc.add_paragraph("")

# Findings
doc.add_heading("Findings & Observations", level=2)
doc.add_paragraph("1. All core functionality works correctly — greeting and counter features behave as coded.")
doc.add_paragraph("2. No input validation on empty name — tapping 'Say Hello' with empty field shows 'Hello, !' (UX issue).")
doc.add_paragraph("3. Counter has no lower bound — can go negative. Worth noting if spec requires non-negative values.")
doc.add_paragraph("4. All UI elements have proper accessibility labels (contentDescription) for testability.")

# Save
output = "/Users/yobel/AndroidStudioProjects/TesAxrail/e2e-test-report.docx"
doc.save(output)
print(f"✅ Report saved to: {output}")
