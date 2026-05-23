"""Generate E2E test report as .docx with embedded screenshot."""
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading("E2E Test Report — TesAxrail App", level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Metadata
doc.add_paragraph("Device: Pixel 7 Emulator (Android 17, emulator-5554)")
doc.add_paragraph("Date: May 21, 2026")
doc.add_paragraph("")

# Test Results
doc.add_heading("Test Results", level=2)

table = doc.add_table(rows=7, cols=5)
table.style = "Table Grid"

# Header row
headers = ["#", "Test Case", "Expected", "Actual", "Status"]
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    cell.paragraphs[0].runs[0].bold = True

# Test data
tests = [
    ("1", "Greeting Flow", '"Hello, Kiro!" appears', '"Hello, Kiro!" appeared', "✅ PASS"),
    ("2", "Counter Increment (x3)", "Counter: 3", "Counter: 3", "✅ PASS"),
    ("3", "Counter Decrement", "Counter: 2", "Counter: 2", "✅ PASS"),
    ("4", "Counter Reset", "Counter: 0", "Counter: 0", "✅ PASS"),
    ("5", "Empty Name Greeting", "Validation or empty", '"Hello, !" displayed', "⚠️ PASS (UX issue)"),
    ("6", "Negative Counter", "Counter: -1", "Counter: -1", "✅ PASS"),
]

for row_idx, test in enumerate(tests, start=1):
    for col_idx, val in enumerate(test):
        table.rows[row_idx].cells[col_idx].text = val

doc.add_paragraph("")

# Screenshot section
doc.add_heading("Screenshot", level=2)
doc.add_paragraph("App state after E2E testing (empty name + counter at -1):")
doc.add_picture(
    "/Users/yobel/AndroidStudioProjects/TesAxrail/e2e-screenshot.png",
    width=Inches(3),
)

# Findings
doc.add_heading("Findings", level=2)
doc.add_paragraph(
    "1. All core functionality works correctly — greeting and counter features behave as coded."
)
doc.add_paragraph(
    '2. UX Issue: Tapping "Say Hello" with an empty name shows "Hello, !" — '
    "the button should be disabled or show validation when the field is empty."
)
doc.add_paragraph(
    "3. No lower bound on counter — counter can go negative. "
    "Worth noting if the spec requires non-negative values only."
)

# Save
output_path = "/Users/yobel/AndroidStudioProjects/TesAxrail/e2e-test-report.docx"
doc.save(output_path)
print(f"Report saved to: {output_path}")
